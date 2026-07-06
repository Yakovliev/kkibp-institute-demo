"""
Convert a finalized Markdown report into a polished Word (.docx) document.

The script uses a two-step workflow:
1. Convert Markdown to DOCX with Pandoc.
2. Post-process the DOCX with python-docx to apply professional report styling.

Usage:
    python tools/markdown_to_professional_docx.py report.md
    python tools/markdown_to_professional_docx.py report.md --output report.docx --title "Report Title"
    python tools/markdown_to_professional_docx.py report.md --update-fields-on-open
    python tools/markdown_to_professional_docx.py report.docx --post-process-only
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DEFAULT_BLUE = (31, 78, 121)
DEFAULT_TEXT = (31, 31, 31)
TABLE_HEADER_FILL = "1F4E79"


def delete_paragraph(paragraph):
    """Remove a paragraph from the document."""
    p = paragraph._element
    p.getparent().remove(p)


def set_style_font(doc, style_name, size=None, bold=None, color=None, font_name="Calibri"):
    """Set font properties on a Word style if that style exists."""
    if style_name not in doc.styles:
        return

    style = doc.styles[style_name]
    font = style.font
    font.name = font_name
    font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=60, start=60, bottom=60, end=60):
    """Set table cell margins in twips."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    """Make a table header row repeat across Word page breaks."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width_pct(table, pct=5000):
    """Set table width as a percentage. Word uses 5000 for 100%."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct))


def add_page_number_footer(doc):
    """Add centered Page {PAGE} footer using a real Word PAGE field."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = "Page "
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def set_word_field_updates_on_open(doc, enabled):
    """Enable or disable Word field updates on open.

    When enabled, Word can show a security/update prompt on open for documents
    containing TOC or other fields. Keep this disabled by default for
    client-facing output unless the caller explicitly wants auto-refresh.
    """
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if enabled:
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")
        return

    if update_fields is not None:
        settings.remove(update_fields)


def clear_dirty_field_flags(doc):
    """Strip w:dirty from every <w:fldChar> in body, headers, and footers.

    Pandoc marks its TOC fldChar dirty, which makes Word display the
    "This document contains fields that may refer to other files. Do you
    want to update the fields in this document?" prompt on every open,
    independently of the document-level w:updateFields setting. Removing the
    per-field dirty flag suppresses the prompt; the TOC can still be
    refreshed manually with Ctrl+A then F9.
    """
    elements = [doc.element]
    for section in doc.sections:
        elements.append(section.header._element)
        elements.append(section.footer._element)

    dirty_attr = qn("w:dirty")
    for element in elements:
        for fld_char in element.iter(qn("w:fldChar")):
            fld_char.attrib.pop(dirty_attr, None)


def clean_front_matter(doc):
    """Split common report front matter into clean centered lines."""
    labels = ["Prepared for:", "Prepared by:", "Report date:", "Testing environment:"]

    for paragraph in doc.paragraphs[:8]:
        text = paragraph.text
        if "Prepared for:" not in text or "Testing environment:" not in text:
            continue

        # Pandoc can collapse Markdown hard line breaks in front matter.
        # Re-split known report metadata labels into separate Word lines.
        normalized = " ".join(text.split())
        pattern = "(" + "|".join(re.escape(label) for label in labels) + ")"
        parts = re.split(pattern, normalized)
        fields = []

        for index in range(1, len(parts), 2):
            label = parts[index]
            value = parts[index + 1].strip() if index + 1 < len(parts) else ""
            next_label_positions = [
                value.find(next_label)
                for next_label in labels
                if next_label in value
            ]
            if next_label_positions:
                value = value[: min(next_label_positions)].strip()
            fields.append(f"{label} {value}".strip())

        if not fields:
            fields = [line.strip() for line in text.splitlines() if line.strip()]

        paragraph.clear()
        for index, line in enumerate(fields):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.size = Pt(10.5)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(14)
        break


def apply_page_setup(doc):
    """Apply compact professional report margins."""
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def remove_duplicate_title(doc):
    """Remove duplicate title when Pandoc metadata title and Markdown H1 match."""
    if len(doc.paragraphs) <= 1:
        return

    first = doc.paragraphs[0]
    second = doc.paragraphs[1]
    if first.text.strip() == second.text.strip() and second.style.name.startswith("Heading"):
        delete_paragraph(second)


def apply_document_styles(doc):
    """Apply report typography to styles and the title paragraph."""
    for style_name in ["Normal", "Body Text", "First Paragraph", "Compact"]:
        set_style_font(doc, style_name, 10.5, False, DEFAULT_TEXT)

    set_style_font(doc, "Title", 20, True, DEFAULT_BLUE)
    set_style_font(doc, "Heading 1", 16, True, DEFAULT_BLUE)
    set_style_font(doc, "Heading 2", 14, True, DEFAULT_BLUE)
    set_style_font(doc, "Heading 3", 12, True, DEFAULT_BLUE)
    set_style_font(doc, "Heading 4", 11, True, DEFAULT_BLUE)

    for style_name in ["Normal", "Body Text", "First Paragraph", "Compact"]:
        if style_name in doc.styles:
            paragraph_format = doc.styles[style_name].paragraph_format
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing = 1.05

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if style_name in doc.styles:
            paragraph_format = doc.styles[style_name].paragraph_format
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
            paragraph_format.keep_with_next = True

    if doc.paragraphs:
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in doc.paragraphs[0].runs:
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*DEFAULT_BLUE)


def is_detailed_results_table(table):
    """Detect the large detailed test-results table used in report appendices."""
    if not table.rows or not table.rows[0].cells:
        return False
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    return headers[:5] == ["No.", "Type", "Title", "Outcome", "Comment"]


def apply_table_styles(doc):
    """Style all tables with professional header and compact body formatting."""
    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_width_pct(table)

        base_size = 7.5 if is_detailed_results_table(table) else 9

        if table.rows:
            repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                shade_cell(cell, TABLE_HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(base_size)

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_margin(cell)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
                        run.font.size = Pt(base_size)


def professionalize_docx(path, update_fields_on_open=False):
    """Apply professional report styling to an existing DOCX in place."""
    doc = Document(path)
    apply_page_setup(doc)
    remove_duplicate_title(doc)
    apply_document_styles(doc)
    clean_front_matter(doc)
    apply_table_styles(doc)
    add_page_number_footer(doc)
    set_word_field_updates_on_open(doc, update_fields_on_open)
    if not update_fields_on_open:
        clear_dirty_field_flags(doc)
    doc.save(path)


def extract_title_from_markdown(markdown_path):
    """Use the first Markdown H1 as default title, falling back to file stem."""
    try:
        for line in markdown_path.read_text(encoding="utf-8").splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
                title = re.sub(r"[*_`]+", "", title)
                return title or markdown_path.stem
    except UnicodeDecodeError:
        pass
    return markdown_path.stem


def default_output_path(input_path):
    """Return input path with .docx suffix."""
    return input_path.with_suffix(".docx")


def build_pandoc_resource_path(markdown_path, extra_paths=None):
    """Build a deterministic Pandoc resource path for relative assets."""
    resource_dirs = []

    for candidate in [markdown_path.parent, *(extra_paths or [])]:
        if not candidate:
            continue

        resolved = Path(candidate).expanduser().resolve()
        if resolved not in resource_dirs and resolved.exists():
            resource_dirs.append(resolved)

    return os.pathsep.join(str(path) for path in resource_dirs)


def run_pandoc(
    markdown_path,
    output_path,
    title,
    from_format="gfm",
    toc=True,
    toc_depth=2,
    reference_doc=None,
    extra_resource_paths=None,
):
    """Run Pandoc to convert Markdown to DOCX."""
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("Pandoc was not found on PATH.")

    command = [
        pandoc,
        str(markdown_path),
        f"--from={from_format}",
        "--to=docx",
        "--standalone",
        "--metadata",
        f"title={title}",
        f"--output={output_path}",
    ]

    resource_path = build_pandoc_resource_path(markdown_path, extra_resource_paths)
    if resource_path:
        command.append(f"--resource-path={resource_path}")

    if toc:
        command.extend(["--toc", f"--toc-depth={toc_depth}"])

    if reference_doc:
        command.append(f"--reference-doc={reference_doc}")

    subprocess.run(command, check=True, cwd=markdown_path.parent)


def validate_docx(path):
    """Validate DOCX package and confirm python-docx can read it."""
    with zipfile.ZipFile(path) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"DOCX zip validation failed at: {bad_member}")
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("DOCX does not contain word/document.xml")

    doc = Document(path)
    return len(doc.paragraphs), len(doc.tables)


def convert_markdown_to_professional_docx(
    input_path,
    output_path=None,
    title=None,
    from_format="gfm",
    toc=True,
    toc_depth=2,
    reference_doc=None,
    update_fields_on_open=False,
):
    """Convert Markdown to DOCX and apply professional styling."""
    invocation_cwd = Path.cwd().resolve()
    input_path = Path(input_path).expanduser().resolve()
    output_path = (
        Path(output_path).expanduser().resolve()
        if output_path
        else default_output_path(input_path)
    )
    reference_doc = Path(reference_doc).expanduser().resolve() if reference_doc else None
    title = title or extract_title_from_markdown(input_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    run_pandoc(
        markdown_path=input_path,
        output_path=output_path,
        title=title,
        from_format=from_format,
        toc=toc,
        toc_depth=toc_depth,
        reference_doc=reference_doc,
        extra_resource_paths=[invocation_cwd],
    )
    professionalize_docx(output_path, update_fields_on_open=update_fields_on_open)
    return output_path


def parse_args(argv):
    parser = argparse.ArgumentParser(
        description="Convert Markdown reports to polished professional DOCX files."
    )
    parser.add_argument(
        "input",
        help="Markdown input path, or DOCX path when --post-process-only is used.",
    )
    parser.add_argument(
        "-o",
        "--output",
        help="Output DOCX path. Defaults to input path with .docx extension.",
    )
    parser.add_argument(
        "--title",
        help="DOCX metadata/title. Defaults to the first Markdown H1 or file name.",
    )
    parser.add_argument(
        "--from-format",
        default="gfm",
        help="Pandoc input format. Default: gfm.",
    )
    parser.add_argument(
        "--no-toc",
        action="store_true",
        help="Do not generate a table of contents.",
    )
    parser.add_argument(
        "--toc-depth",
        type=int,
        default=2,
        help="Pandoc TOC depth. Default: 2.",
    )
    parser.add_argument(
        "--reference-doc",
        help="Optional Pandoc reference DOCX template.",
    )
    parser.add_argument(
        "--post-process-only",
        action="store_true",
        help="Treat input as an existing DOCX and only apply professional styling.",
    )
    parser.add_argument(
        "--skip-validation",
        action="store_true",
        help="Skip DOCX structural validation after generation.",
    )
    parser.add_argument(
        "--update-fields-on-open",
        action="store_true",
        help=(
            "Ask Word to auto-update fields such as TOC on open. "
            "Disabled by default because it can trigger a Word prompt."
        ),
    )
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv or sys.argv[1:])
    input_path = Path(args.input)

    if args.post_process_only:
        if input_path.suffix.lower() != ".docx":
            raise SystemExit("--post-process-only requires a .docx input file.")
        output_path = Path(args.output) if args.output else input_path
        if output_path != input_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
        professionalize_docx(
            output_path,
            update_fields_on_open=args.update_fields_on_open,
        )
    else:
        if input_path.suffix.lower() not in [".md", ".markdown"]:
            raise SystemExit("Input must be a Markdown file unless --post-process-only is used.")
        output_path = convert_markdown_to_professional_docx(
            input_path=input_path,
            output_path=args.output,
            title=args.title,
            from_format=args.from_format,
            toc=not args.no_toc,
            toc_depth=args.toc_depth,
            reference_doc=args.reference_doc,
            update_fields_on_open=args.update_fields_on_open,
        )

    if not args.skip_validation:
        paragraph_count, table_count = validate_docx(output_path)
        print(f"Validated DOCX: {output_path}")
        print(f"Paragraphs: {paragraph_count}; tables: {table_count}")
    else:
        print(f"Wrote DOCX: {output_path}")


if __name__ == "__main__":
    main()
